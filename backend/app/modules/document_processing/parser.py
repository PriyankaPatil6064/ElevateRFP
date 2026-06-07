# app/modules/document_processing/parser.py
from typing import Dict, List, Optional, Tuple
import os
import tempfile
from pathlib import Path
import structlog
import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.config import settings
from app.modules.document_processing.ocr import OCRProcessor

logger = structlog.get_logger()

class DocumentProcessor:
    """Enterprise document processing with OCR and intelligent extraction"""
    
    def __init__(self):
        self.ocr_processor = OCRProcessor()
        self.requirement_extractor = None  # extractor.py not yet implemented
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    async def process_document(self, file_path: str, filename: str) -> Dict:
        """Process uploaded document and extract structured information"""
        try:
            logger.info("Starting document processing", filename=filename)
            
            # Determine file type
            file_extension = Path(filename).suffix.lower()
            
            # Extract raw text
            if file_extension == '.pdf':
                raw_text, metadata = await self._process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                raw_text, metadata = await self._process_docx(file_path)
            elif file_extension == '.txt':
                raw_text, metadata = await self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Clean and preprocess text
            processed_text = self._clean_text(raw_text)
            
            # Extract structured requirements
            requirements = await self.requirement_extractor.extract_requirements(processed_text) if self.requirement_extractor else []
            
            # Create document chunks for vector storage
            chunks = self._create_chunks(processed_text, metadata)
            
            # Extract metadata
            document_metadata = self._extract_document_metadata(processed_text, filename)
            
            result = {
                "raw_text": raw_text,
                "processed_text": processed_text,
                "requirements": requirements,
                "chunks": chunks,
                "metadata": {**metadata, **document_metadata},
                "file_info": {
                    "filename": filename,
                    "extension": file_extension,
                    "size": os.path.getsize(file_path)
                }
            }
            
            logger.info("Document processing completed", 
                       filename=filename, 
                       text_length=len(processed_text),
                       requirements_count=len(requirements),
                       chunks_count=len(chunks))
            
            return result
            
        except Exception as e:
            logger.error("Document processing failed", filename=filename, error=str(e))
            raise
    
    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Process PDF document with OCR fallback"""
        text_content = []
        metadata = {"pages": 0, "has_images": False, "ocr_used": False}
        
        try:
            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)
            metadata["pages"] = len(doc)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try to extract text directly
                page_text = page.get_text()
                
                # If no text or very little text, use OCR
                if len(page_text.strip()) < 50:
                    logger.info("Using OCR for page", page_num=page_num)
                    
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img_data = pix.tobytes("png")
                    
                    # Save to temporary file for OCR
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                        temp_img.write(img_data)
                        temp_img_path = temp_img.name
                    
                    try:
                        # Perform OCR
                        ocr_text = await self.ocr_processor.extract_text(temp_img_path)
                        page_text = ocr_text
                        metadata["ocr_used"] = True
                    finally:
                        os.unlink(temp_img_path)
                
                text_content.append(page_text)
                
                # Check for images
                if page.get_images():
                    metadata["has_images"] = True
            
            doc.close()
            
        except Exception as e:
            logger.error("PDF processing failed", error=str(e))
            raise
        
        return "\n\n".join(text_content), metadata
    
    async def _process_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Process DOCX document"""
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "has_images": len(doc.inline_shapes) > 0
            }
            
            return "\n\n".join(text_content), metadata
            
        except Exception as e:
            logger.error("DOCX processing failed", error=str(e))
            raise
    
    async def _process_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Process plain text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                "lines": len(content.split('\n')),
                "encoding": "utf-8"
            }
            
            return content, metadata
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    metadata = {
                        "lines": len(content.split('\n')),
                        "encoding": encoding
                    }
                    return content, metadata
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Remove excessive spaces
                line = ' '.join(line.split())
                cleaned_lines.append(line)
        
        # Join lines and normalize spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove page numbers and headers/footers (basic patterns)
        import re
        cleaned_text = re.sub(r'\n\s*Page \d+.*?\n', '\n', cleaned_text, flags=re.IGNORECASE)
        cleaned_text = re.sub(r'\n\s*\d+\s*\n', '\n', cleaned_text)
        
        return cleaned_text
    
    def _create_chunks(self, text: str, metadata: Dict) -> List[Document]:
        """Create document chunks for vector storage"""
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    **metadata,
                    "chunk_id": i,
                    "chunk_size": len(chunk)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _extract_document_metadata(self, text: str, filename: str) -> Dict:
        """Extract metadata from document content"""
        import re
        from datetime import datetime
        
        metadata = {}
        
        # Extract dates
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b',
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        ]
        
        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))
        
        if dates:
            metadata["extracted_dates"] = dates[:5]  # Limit to first 5 dates
        
        # Extract email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            metadata["contact_emails"] = list(set(emails))[:3]  # Limit and deduplicate
        
        # Extract phone numbers
        phone_pattern = r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b'
        phones = re.findall(phone_pattern, text)
        if phones:
            metadata["contact_phones"] = list(set(phones))[:3]
        
        # Estimate document complexity
        word_count = len(text.split())
        sentence_count = len(re.findall(r'[.!?]+', text))
        
        metadata.update({
            "word_count": word_count,
            "sentence_count": sentence_count,
            "avg_sentence_length": word_count / max(sentence_count, 1),
            "complexity_score": min(100, (word_count / 100) + (sentence_count / 10))
        })
        
        return metadata